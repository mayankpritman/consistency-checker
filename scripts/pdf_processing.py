# pdf_processing.py
import PyPDF2
import pdfplumber
from docx import Document

# Function to extract text from PDF
#def extract_text_from_pdf(file):
#    reader = PyPDF2.PdfReader(file)
#    text = ''
#    for page in range(len(reader.pages)):
#        text += reader.pages[page].extract_text()
#    return text

#############changes done#######################
def normalize_text(text):
    return ' '.join(text.strip().split()).lower()
def extract_text_from_pdf(file):
    text = ""
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text
    return text

# Function to extract tables from PDF
def extract_tables_from_pdf(filepath):
    try:
        with pdfplumber.open(filepath) as pdf:
            tables = []
            for page in pdf.pages:
                extracted_tables = page.extract_tables()
                for table in extracted_tables:
                    tables.append(table)
        return tables
    except Exception as e:
        print(f"An error occurred: {e}")
        return []

# Function to extract text from DOCX
def extract_text_from_docx(filepath):
    doc = Document(filepath)
    text = []
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)
    return '\n'.join(text)

# Function to extract tables from DOCX
def extract_tables_from_docx(filepath):
    doc = Document(filepath)
    tables = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text)
            table_data.append(row_data)
        tables.append(table_data)
    return tables

# Function to extract text from TXT
def extract_text_from_txt(filepath):
    with open(filepath, 'r', encoding='utf-8') as file:
        text = file.read()
    return text